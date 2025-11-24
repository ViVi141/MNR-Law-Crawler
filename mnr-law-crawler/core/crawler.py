"""
爬虫核心模块 - 政策爬取主逻辑
适配自然资源部政府信息公开平台
"""

import os
import json
import time
import logging
from typing import Dict, List, Optional, Callable, Any
from datetime import datetime
from urllib.parse import urljoin
from bs4 import BeautifulSoup

from .config import Config
from .api_client import APIClient
from .converter import DocumentConverter
from .models import Policy, PolicyDetail, FileAttachment, CrawlProgress


# 自然资源部分类配置
MNR_CATEGORIES = {
    '自然资源调查监测': {'code': '1318', 'name': '自然资源调查监测'},
    '自然资源确权登记': {'code': '1319', 'name': '自然资源确权登记'},
    '自然资源合理开发利用': {'code': '1320', 'name': '自然资源合理开发利用'},
    '自然资源有偿使用': {'code': '1321', 'name': '自然资源有偿使用'},
    '国土空间规划': {'code': '1322', 'name': '国土空间规划'},
    '国土空间用途管制': {'code': '1663', 'name': '国土空间用途管制'},
    '国土空间生态修复': {'code': '1324', 'name': '国土空间生态修复'},
    '耕地保护': {'code': '1325', 'name': '耕地保护'},
    '地质勘查': {'code': '1326', 'name': '地质勘查'},
    '矿产勘查': {'code': '1327', 'name': '矿产勘查'},
    '矿产保护': {'code': '1328', 'name': '矿产保护'},
    '矿产开发': {'code': '1329', 'name': '矿产开发'},
    '地质环境保护': {'code': '1330', 'name': '地质环境保护'},
    '海洋资源': {'code': '1331', 'name': '海洋资源'},
    '测绘地理信息': {'code': '1332', 'name': '测绘地理信息'},
    '地质灾害防治': {'code': '1334', 'name': '地质灾害防治'},
    '地质公园': {'code': '1335', 'name': '地质公园'},
    '地质遗迹保护': {'code': '1336', 'name': '地质遗迹保护'},
    '矿业权评估': {'code': '1338', 'name': '矿业权评估'},
    '机构建设': {'code': '1339', 'name': '机构建设'},
    '综合管理': {'code': '1340', 'name': '综合管理'},
    '其他': {'code': '1341', 'name': '其他'}
}


class PolicyCrawler:
    """政策爬虫核心类 - 适配自然资源部API"""
    
    def __init__(self, config: Config, progress_callback: Optional[Callable] = None):
        """初始化爬虫
        
        Args:
            config: 配置对象
            progress_callback: 进度回调函数 (progress: CrawlProgress) -> None
        """
        self.config = config
        self.api_client = APIClient(config)
        self.converter = DocumentConverter()
        self.progress_callback = progress_callback
        self.stop_requested = False  # 停止标志
        self.progress = CrawlProgress()
        self.base_url = config.get("base_url", "https://gi.mnr.gov.cn/")
        
        # 创建输出目录
        self._create_output_dirs()
    
    def _create_output_dirs(self):
        """创建输出目录"""
        output_dir = self.config.output_dir
        os.makedirs(output_dir, exist_ok=True)
        os.makedirs(f"{output_dir}/json", exist_ok=True)
        os.makedirs(f"{output_dir}/files", exist_ok=True)
        os.makedirs(f"{output_dir}/markdown", exist_ok=True)
        os.makedirs(f"{output_dir}/docx", exist_ok=True)
    
    def request_stop(self):
        """请求停止爬取"""
        self.stop_requested = True
        logging.info("\n[停止] 收到停止请求，正在停止...")
    
    def _update_progress(self, **kwargs):
        """更新进度并触发回调"""
        for key, value in kwargs.items():
            setattr(self.progress, key, value)
        
        if self.progress_callback:
            self.progress_callback(self.progress)
    
    def _parse_date(self, date_str: str) -> Optional[datetime]:
        """解析日期字符串为datetime对象"""
        if not date_str:
            return None
        for fmt in ('%Y年%m月%d日', '%Y-%m-%d', '%Y/%m/%d', '%Y.%m.%d'):
            try:
                return datetime.strptime(date_str, fmt)
            except Exception:
                continue
        return None
    
    def _parse_json_results(self, data: Dict, callback: Optional[Callable] = None) -> List[Policy]:
        """解析JSON格式的搜索结果"""
        policies = []
        
        try:
            # 根据实际返回的JSON结构解析
            if 'results' in data:
                items = data['results']
            elif 'data' in data:
                items = data['data']
            elif isinstance(data, list):
                items = data
            else:
                items = []
            
            for item in items:
                # 尝试获取内容（如果API返回）
                content = item.get('content', '').strip()
                # 如果API没有返回内容，尝试从其他字段获取
                if not content:
                    content = item.get('summary', '').strip() or item.get('abstract', '').strip()
                
                # 解析并格式化日期
                raw_date = item.get('pubdate', item.get('publishdate', ''))
                pub_date = ''
                if raw_date:
                    parsed_date = self._parse_date(raw_date)
                    if parsed_date:
                        pub_date = parsed_date.strftime('%Y-%m-%d')
                    else:
                        # 如果解析失败，尝试直接使用原始值（可能已经是标准格式）
                        pub_date = raw_date.strip()
                
                policy = Policy(
                    title=item.get('title', '') or '',
                    pub_date=pub_date or '',
                    doc_number=item.get('filenum', '') or '',
                    source=item.get('url', '') or '',
                    link=item.get('url', '') or '',
                    url=item.get('url', '') or '',
                    content=content or '',
                    category=item.get('category', '') or '',
                    validity=item.get('status', '') or '',
                    effective_date=item.get('effectivedate', '') or '',
                    crawl_time=datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                )
                policies.append(policy)
                    
        except Exception as e:
            if callback:
                callback(f"解析JSON结果失败: {e}")
        
        return policies
    
    def _parse_html_results(self, soup: BeautifulSoup, callback: Optional[Callable] = None, category_name: str = '', max_policies: Optional[int] = None, data_source: Optional[Dict[str, Any]] = None) -> List[Policy]:
        """解析HTML格式的搜索结果
        
        Args:
            soup: BeautifulSoup对象
            callback: 进度回调函数
            category_name: 分类名称
            max_policies: 最大解析数量（None表示不限制）
        """
        policies = []
        
        try:
            # 查找政策列表 - 适配新的网站结构
            table = soup.find('table', class_='table')
            if not table:
                if callback:
                    callback("未找到政策列表表格")
                return policies
            
            # 查找所有政策行（跳过表头）
            rows = table.find_all('tr')[1:]  # 跳过第一行表头
            total_rows = len(rows)
            if callback:
                callback(f"HTML表格中找到 {total_rows} 行数据")
            
            # 如果设置了最大数量，只处理前N条有效政策
            parsed_count = 0
            
            for row in rows:
                # 如果已达到最大数量，停止解析
                if max_policies is not None and parsed_count >= max_policies:
                    break
                try:
                    # 获取所有单元格
                    cells = row.find_all('td')
                    if len(cells) < 4:
                        continue
                    
                    # 检查是否是主政策行（不是详细信息行）
                    first_cell = cells[0].get_text(strip=True)
                    if not first_cell or first_cell in ['标    题', '索    引', '发文字号', '生成日期', '实施日期']:
                        # 这是详细信息行，跳过
                        continue
                    
                    # 检查是否是有效的政策索引号（应该包含年份和编号）
                    if not first_cell or len(first_cell) < 4 or not first_cell[0].isdigit():
                        continue
                    
                    # 解析表格数据 - 适配新结构
                    title_cell = cells[1]
                    doc_number = cells[2].get_text(strip=True)
                    pub_date = cells[3].get_text(strip=True)
                    
                    # 获取标题和链接 - 新网站的结构
                    title_link = title_cell.find('a', target='_blank')
                    if not title_link:
                        # 尝试其他方式查找链接
                        title_link = title_cell.find('a')
                    
                    if not title_link:
                        continue
                    
                    title = title_link.get_text(strip=True)
                    detail_url = title_link.get('href', '')
                    
                    # 检查解析的数据是否合理
                    if doc_number == "标    题" or pub_date == title:
                        # 如果发文字号是"标    题"或发布日期是标题内容，说明解析到了详细信息
                        # 尝试从详细信息中获取正确的数据
                        detail_info = title_cell.find('div', class_='box')
                        if detail_info:
                            detail_table = detail_info.find('table')
                            if detail_table:
                                detail_rows = detail_table.find_all('tr')
                                for detail_row in detail_rows:
                                    detail_cells = detail_row.find_all('td')
                                    if len(detail_cells) >= 2:
                                        label = detail_cells[0].get_text(strip=True)
                                        value = detail_cells[1].get_text(strip=True)
                                        
                                        if '发文字号' in label:
                                            doc_number = value
                                        elif '生成日期' in label:
                                            pub_date = value
                    
                    # 构建完整链接
                    if detail_url and not detail_url.startswith('http'):
                        detail_url = urljoin(self.base_url, detail_url)
                    
                    # 解析并格式化日期
                    pub_date_formatted = ''
                    if pub_date:
                        parsed_date = self._parse_date(pub_date)
                        if parsed_date:
                            pub_date_formatted = parsed_date.strftime('%Y-%m-%d')
                        else:
                            # 如果解析失败，尝试直接使用原始值（可能已经是标准格式）
                            pub_date_formatted = pub_date.strip()
                    
                    policy = Policy(
                        title=title,
                        pub_date=pub_date_formatted,
                        doc_number=doc_number,
                        source=detail_url,
                        link=detail_url,
                        url=detail_url,
                        content='',  # 初始为空，会在crawl_policies中填充
                        category=category_name,
                        crawl_time=datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    )
                    # 保存数据源信息到policy对象（通过添加自定义属性）
                    policy._data_source = data_source
                    
                    policies.append(policy)
                    parsed_count += 1  # 成功解析一条政策
                        
                except Exception as e:
                    if callback:
                        callback(f"解析政策项失败: {e}")
                    continue
            
            # 统计信息
            if callback:
                callback(f"成功解析 {len(policies)} 条政策（从 {total_rows} 行中）")
                    
        except Exception as e:
            if callback:
                callback(f"解析HTML结果失败: {e}")
        
        return policies
    
    def search_all_policies(
        self,
        keywords: List[str] = None,
        start_date: str = None,
        end_date: str = None,
        callback: Optional[Callable] = None,
        limit_pages: Optional[int] = None
    ) -> List[Policy]:
        """搜索所有政策（支持多个数据源）
        
        Args:
            keywords: 关键词列表
            start_date: 起始日期 yyyy-MM-dd
            end_date: 结束日期 yyyy-MM-dd
            callback: 进度回调函数
            limit_pages: 限制最大页数（用于测试模式，None表示不限制）
            
        Returns:
            政策列表
        """
        if keywords is None:
            keywords = []
        
        # 获取启用的数据源列表
        data_sources = self.config.get("data_sources", [])
        if not data_sources:
            # 如果没有配置数据源列表，使用默认配置（向后兼容）
            data_sources = [{
                "name": "默认数据源",
                "base_url": self.config.get("base_url", "https://gi.mnr.gov.cn/"),
                "search_api": self.config.get("search_api", "https://search.mnr.gov.cn/was5/web/search"),
                "channel_id": self.config.get("channel_id", "216640"),
                "enabled": True
            }]
        
        # 过滤出启用的数据源
        enabled_sources = [ds for ds in data_sources if ds.get("enabled", True)]
        
        if callback:
            callback(f"开始搜索政策，关键词: {', '.join(keywords) if keywords else '(无关键词，搜索全部政策)'}")
            callback(f"启用数据源: {', '.join([ds.get('name', '未知') for ds in enabled_sources])}")
        
        all_policies = []
        seen_ids = set()  # 用于去重（跨数据源）
        
        # 遍历每个数据源
        for data_source in enabled_sources:
            source_name = data_source.get("name", "未知数据源")
            if callback:
                callback(f"\n{'='*60}")
                callback(f"开始爬取数据源: {source_name}")
                callback(f"{'='*60}")
            
            policies = []
            page = 1
            if limit_pages is not None:
                max_pages = limit_pages
            else:
                max_pages = self.config.get("max_pages", 999999)
            max_empty_pages = self.config.get("max_empty_pages", 3)
            consecutive_empty_pages = 0
            
            while page <= max_pages:
                if self.stop_requested:
                    if callback:
                        callback("停止搜索")
                    break
            
            if callback:
                callback(f"正在抓取第{page}页...")
            
            try:
                result = self.api_client.search_policies(keywords, page, start_date, end_date, data_source)
                
                # 解析结果
                page_policies = []
                requested_perpage = self.config.get("perpage", 20)
                
                if not result:
                    # API请求失败或无响应
                    consecutive_empty_pages += 1
                    if callback:
                        callback(f"第{page}页API请求失败或无响应")
                else:
                    # 解析响应数据
                    if result['type'] == 'json':
                        page_policies = self._parse_json_results(result['data'], callback)
                        if callback:
                            callback(f"JSON解析: 请求{requested_perpage}条，实际返回{len(page_policies)}条")
                    elif result['type'] == 'html':
                        soup = BeautifulSoup(result['data'], 'html.parser')
                        # 限制解析数量为请求的perpage值（因为API返回的行数可能是perpage的倍数）
                        page_policies = self._parse_html_results(soup, callback, '全部', max_policies=requested_perpage, data_source=data_source)
                        if callback:
                            callback(f"HTML解析: 请求{requested_perpage}条，实际解析{len(page_policies)}条")
                
                # 检查是否解析出政策
                if not page_policies or len(page_policies) == 0:
                    consecutive_empty_pages += 1
                    if callback:
                        callback(f"第{page}页无数据（解析出0条政策），连续空页数: {consecutive_empty_pages}/{max_empty_pages}")
                    
                    # 检查是否达到连续空页限制
                    if consecutive_empty_pages >= max_empty_pages:
                        if callback:
                            callback(f"连续{max_empty_pages}页无数据，停止爬取")
                        break
                    
                    # 继续下一页
                    page += 1
                    continue
                else:
                    # 有数据，重置连续空页计数
                    consecutive_empty_pages = 0
                
                # 去重并添加到结果
                new_policies_count = 0
                for policy in page_policies:
                    policy_id = policy.id
                    if policy_id not in seen_ids:
                        seen_ids.add(policy_id)
                        policies.append(policy)
                        new_policies_count += 1
                
                if callback:
                    callback(f"第{page}页获取{len(page_policies)}条政策（新增{new_policies_count}条），累计{len(policies)}条")
                
                # 如果连续多页都没有新政策（全部重复），可能已经爬取完毕
                if new_policies_count == 0:
                    consecutive_empty_pages += 1
                    if consecutive_empty_pages >= max_empty_pages:
                        if callback:
                            callback(f"连续{max_empty_pages}页无新政策（全部重复），停止爬取")
                        break
                else:
                    consecutive_empty_pages = 0  # 有新政策，重置计数
                
                # 控制速度
                time.sleep(self.config.get("request_delay", 2))
                
                page += 1
                
            except Exception as e:
                if callback:
                    callback(f"第{page}页抓取失败: {e}")
                break
            
            # 合并当前数据源的政策到总列表（去重）
            for policy in policies:
                policy_id = policy.id
                if policy_id not in seen_ids:
                    seen_ids.add(policy_id)
                    all_policies.append(policy)
            
            if callback:
                callback(f"数据源 {source_name} 完成，共获取 {len(policies)} 条政策")
        
        if callback:
            callback(f"\n所有数据源爬取完成，总计 {len(all_policies)} 条政策")
        
        return all_policies
    
    def crawl_single_policy(self, policy: Policy, callback: Optional[Callable] = None) -> bool:
        """爬取单个政策
        
        Args:
            policy: 政策对象
            callback: 进度回调函数
            
        Returns:
            是否成功
        """
        self._update_progress(
            current_policy_id=policy.id,
            current_policy_title=policy.title
        )
        
        if callback:
            callback(f"\n爬取政策: {policy.title}")
        
        logging.info(f"\n爬取政策: {policy.title}")
        logging.info("=" * 60)
        
        # 1. 获取详情页内容和附件（如果还没有）
        attachments = []
        if not policy.content and policy.link:
            if callback:
                callback("获取详情页内容和附件...")
            # 获取数据源信息（如果policy对象有保存）
            data_source = getattr(policy, '_data_source', None)
            detail_result = self.api_client.get_policy_detail(policy.link, data_source)
            policy.content = detail_result.get('content', '')
            attachments = detail_result.get('attachments', [])
        
        # 2. 保存JSON数据
        if self.config.get("save_json", True):
            self._save_json(policy)
        
        # 3. 获取文件编号（markdown 和 files 文件夹各自独立递增）
        markdown_number = self._get_next_markdown_number()
        file_number = self._get_next_file_number()
        
        # 4. 下载附件（如果启用）
        if self.config.get("save_files", True) and attachments:
            self._download_attachments(policy, attachments, file_number, callback)
        
        # 5. 生成RAG Markdown
        if self.config.get("save_markdown", True):
            self._generate_rag_markdown(policy, markdown_number)
        
        # 6. 生成DOCX格式（如果启用）
        if self.config.get("save_docx", True):
            self._generate_docx(policy, markdown_number, callback)
        
        if callback:
            callback("   ✓ 政策详细内容爬取完成")
        logging.info("   ✓ 政策详细内容爬取完成")
        return True
    
    def _save_json(self, policy: Policy):
        """保存JSON数据"""
        policy_id = policy.id.replace('|', '_').replace('/', '_')[:50]  # 简化ID作为文件名
        filepath = f"{self.config.output_dir}/json/policy_{policy_id}.json"
        
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(policy.to_dict(), f, ensure_ascii=False, indent=2)
            logging.info(f"[OK] JSON已保存: {filepath}")
        except Exception as e:
            logging.info(f"[X] JSON保存失败: {e}")
    
    def _generate_rag_markdown(
        self,
        policy: Policy,
        file_number: int
    ):
        """生成RAG格式的Markdown文件"""
        try:
            md_lines = []
            
            # YAML Front Matter
            md_lines.append('---')
            md_lines.append(f'title: "{policy.title}"')
            md_lines.append(f'level: "{policy.level}"')
            md_lines.append(f'category: "{policy.category}"')
            md_lines.append(f'pub_date: "{policy.pub_date}"')
            md_lines.append(f'doc_number: "{policy.doc_number}"')
            if policy.effective_date:
                md_lines.append(f'effective_date: "{policy.effective_date}"')
            if policy.validity:
                md_lines.append(f'validity: "{policy.validity}"')
            md_lines.append(f'source_url: "{policy.source}"')
            md_lines.append(f'crawl_time: "{policy.crawl_time}"')
            md_lines.append('---')
            md_lines.append('')
            
            md_lines.append(f'# {policy.title}')
            md_lines.append('')
            
            md_lines.append('## 基本信息')
            md_lines.append('')
            md_lines.append(f'- **发布机构**: {policy.level}')
            md_lines.append(f'- **发布日期**: {policy.pub_date}')
            if policy.doc_number:
                md_lines.append(f'- **发文字号**: {policy.doc_number}')
            if policy.effective_date:
                md_lines.append(f'- **生效日期**: {policy.effective_date}')
            if policy.validity:
                md_lines.append(f'- **有效性**: {policy.validity}')
            if policy.category:
                md_lines.append(f'- **分类**: {policy.category}')
            md_lines.append(f'- **来源链接**: [查看原文]({policy.source})')
            md_lines.append('')
            
            md_lines.append('---')
            md_lines.append('')
            md_lines.append('## 正文内容')
            md_lines.append('')
            if policy.content:
                md_lines.append(policy.content)
            else:
                md_lines.append('> **注意**: 该政策的正文内容无法自动获取。')
                md_lines.append('> ')
                md_lines.append('> 请访问[来源链接](#基本信息)查看完整文档内容。')
            
            safe_title = "".join(c for c in policy.title if c.isalnum() or c in (' ', '-', '_')).strip()
            if not safe_title:
                safe_title = f"政策_{policy.id[:8]}"
            
            md_filename = f"{file_number:04d}_{safe_title}.md"
            md_filepath = f"{self.config.output_dir}/markdown/{md_filename}"
            
            with open(md_filepath, 'w', encoding='utf-8') as f:
                f.write('\n'.join(md_lines))
            
            logging.info(f"[OK] Markdown已保存: {md_filepath}")
            
        except Exception as e:
            logging.info(f"[X] Markdown生成失败: {e}")
    
    def _generate_docx(
        self,
        policy: Policy,
        file_number: int,
        callback: Optional[Callable] = None
    ):
        """生成DOCX格式文件
        
        Args:
            policy: 政策对象
            file_number: 文件编号
            callback: 进度回调函数
        """
        try:
            # 尝试导入 python-docx
            try:
                from docx import Document
                from docx.shared import Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError:
                if callback:
                    callback("  [X] python-docx未安装，无法生成DOCX")
                logging.info("[X] python-docx未安装，无法生成DOCX")
                return
            
            # 创建文档
            doc = Document()
            
            # 设置文档标题
            title = doc.add_heading(policy.title, level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加基本信息
            doc.add_heading('基本信息', level=2)
            info_para = doc.add_paragraph()
            info_para.add_run('发布机构: ').bold = True
            info_para.add_run(policy.level)
            info_para.add_run('\n发布日期: ').bold = True
            info_para.add_run(policy.pub_date)
            
            if policy.doc_number:
                info_para.add_run('\n发文字号: ').bold = True
                info_para.add_run(policy.doc_number)
            
            if policy.effective_date:
                info_para.add_run('\n生效日期: ').bold = True
                info_para.add_run(policy.effective_date)
            
            if policy.validity:
                info_para.add_run('\n有效性: ').bold = True
                info_para.add_run(policy.validity)
            
            if policy.category:
                info_para.add_run('\n分类: ').bold = True
                info_para.add_run(policy.category)
            
            info_para.add_run('\n来源链接: ').bold = True
            info_para.add_run(policy.source)
            
            # 添加正文内容
            doc.add_heading('正文内容', level=2)
            
            if policy.content:
                # 将正文内容按段落分割并添加
                content_lines = policy.content.split('\n')
                for line in content_lines:
                    line = line.strip()
                    if line:
                        # 检查是否是标题（以#开头）
                        if line.startswith('#'):
                            # 计算标题级别
                            level = len(line) - len(line.lstrip('#'))
                            heading_text = line.lstrip('#').strip()
                            if heading_text:
                                doc.add_heading(heading_text, level=min(level, 9))
                        else:
                            doc.add_paragraph(line)
            else:
                doc.add_paragraph('该政策的正文内容无法自动获取。')
                doc.add_paragraph(f'请访问来源链接查看完整文档内容: {policy.source}')
            
            # 保存文件
            safe_title = "".join(c for c in policy.title if c.isalnum() or c in (' ', '-', '_')).strip()
            if not safe_title:
                safe_title = f"政策_{policy.id[:8]}"
            
            docx_filename = f"{file_number:04d}_{safe_title}.docx"
            docx_filepath = f"{self.config.output_dir}/docx/{docx_filename}"
            
            doc.save(docx_filepath)
            
            if callback:
                callback(f"  [OK] DOCX已保存: {docx_filename}")
            logging.info(f"[OK] DOCX已保存: {docx_filepath}")
            
        except Exception as e:
            if callback:
                callback(f"  [X] DOCX生成失败: {e}")
            logging.info(f"[X] DOCX生成失败: {e}")
    
    def _get_next_markdown_number(self) -> int:
        """获取下一个 Markdown 文件编号"""
        markdown_dir = f"{self.config.output_dir}/markdown"
        
        if not os.path.exists(markdown_dir):
            return 1
        
        existing_files = [f for f in os.listdir(markdown_dir) if f.endswith('.md')]
        
        if not existing_files:
            return 1
        
        numbers = []
        for filename in existing_files:
            parts = filename.split('_', 1)
            if parts and len(parts) >= 2 and parts[0].isdigit():
                numbers.append(int(parts[0]))
        
        if not numbers:
            return 1
        
        return max(numbers) + 1
    
    def _get_next_file_number(self) -> int:
        """获取下一个附件文件编号"""
        files_dir = f"{self.config.output_dir}/files"
        
        if not os.path.exists(files_dir):
            return 1
        
        existing_files = os.listdir(files_dir)
        
        if not existing_files:
            return 1
        
        numbers = []
        for filename in existing_files:
            parts = filename.split('_', 1)
            if parts and len(parts) >= 2 and parts[0].isdigit():
                numbers.append(int(parts[0]))
        
        if not numbers:
            return 1
        
        return max(numbers) + 1
    
    def _download_attachments(
        self,
        policy: Policy,
        attachments: List[Dict[str, str]],
        file_number: int,
        callback: Optional[Callable] = None
    ):
        """下载附件
        
        Args:
            policy: 政策对象
            attachments: 附件列表，每个附件包含 {'url': str, 'name': str}
            file_number: 文件编号
            callback: 进度回调函数
        """
        if not attachments:
            return
        
        # 筛选需要下载的文件
        target_files = []
        
        if self.config.get("download_all_files", False):
            target_files = attachments
            if callback:
                callback(f"\n[下载所有文件] 已启用，将下载所有 {len(attachments)} 个附件")
        else:
            # 根据文件扩展名筛选
            for att in attachments:
                url = att.get('url', '')
                name = att.get('name', '')
                
                # 从URL或名称中提取扩展名
                file_ext = ''
                if url:
                    file_ext = os.path.splitext(url.split('?')[0])[1].lower().strip('.')
                if not file_ext and name:
                    file_ext = os.path.splitext(name)[1].lower().strip('.')
                
                should_download = False
                if file_ext == 'docx' and self.config.get("download_docx", True):
                    should_download = True
                elif file_ext == 'doc' and self.config.get("download_doc", True):
                    should_download = True
                elif file_ext == 'pdf' and self.config.get("download_pdf", False):
                    should_download = True
                elif file_ext in ['zip', 'tar', 'rar', '7z', 'gz']:
                    # 压缩文件总是下载（包括.tar, .zip等）
                    should_download = True
                
                if should_download:
                    target_files.append(att)
            
            if target_files:
                if callback:
                    callback(f"\n从 {len(attachments)} 个附件中筛选出 {len(target_files)} 个文件")
        
        if not target_files:
            if callback:
                callback("没有需要下载的附件")
            return
        
        # 下载文件
        safe_title = "".join(c for c in policy.title if c.isalnum() or c in (' ', '-', '_')).strip()
        if not safe_title:
            safe_title = f"政策_{policy.id[:8]}"
        
        for i, att in enumerate(target_files, 1):
            url = att.get('url', '')
            name = att.get('name', '')
            
            if callback:
                callback(f"\n  [{i}/{len(target_files)}] 下载: {name or url}")
            
            # 确定保存文件名
            if name:
                # 清理文件名
                safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_', '.')).strip()
                if not safe_name:
                    safe_name = f"附件_{i}"
                # 确保有扩展名
                if not os.path.splitext(safe_name)[1]:
                    # 从URL提取扩展名
                    ext = os.path.splitext(url.split('?')[0])[1]
                    if ext:
                        safe_name += ext
            else:
                # 从URL提取文件名
                safe_name = url.split('/')[-1].split('?')[0]
                if not safe_name or len(safe_name) < 3:
                    safe_name = f"附件_{i}"
            
            # 如果只有一个附件且名称与政策标题相同，使用简化名称
            if len(target_files) == 1 and safe_name == policy.title:
                ext = os.path.splitext(safe_name)[1] or os.path.splitext(url.split('?')[0])[1]
                save_filename = f"{file_number:04d}_{safe_title}{ext}"
            else:
                save_filename = f"{file_number:04d}_{safe_title}_{safe_name}"
            
            save_path = f"{self.config.output_dir}/files/{save_filename}"
            
            # 下载文件
            if self.api_client.download_file(url, save_path):
                if callback:
                    callback(f"    [OK] 下载成功: {save_filename}")
                logging.info(f"[OK] 附件下载成功: {save_path}")
            else:
                if callback:
                    callback(f"    [X] 下载失败: {name or url}")
                logging.info(f"[X] 附件下载失败: {url}")
            
            # 下载间隔
            if i < len(target_files):
                time.sleep(0.3)
    
    def crawl_batch(
        self,
        keywords: List[str] = None,
        start_date: str = None,
        end_date: str = None,
        callback: Optional[Callable] = None
    ) -> CrawlProgress:
        """批量爬取"""
        if not hasattr(self, 'progress') or self.progress is None:
            self.progress = CrawlProgress()
            self.progress.start_time = datetime.now()
        elif self.progress.start_time is None:
            self.progress.start_time = datetime.now()
        else:
            original_start_time = self.progress.start_time
            self.progress = CrawlProgress()
            self.progress.start_time = original_start_time
        
        self._update_progress()
        
        # 1. 搜索所有政策
        all_policies = self.search_all_policies(keywords, start_date, end_date, callback)
        
        if self.stop_requested:
            self.progress.end_time = datetime.now()
            self._update_progress()
            return self.progress
        
        self.progress.total_count = len(all_policies)
        self._update_progress()
        
        if callback:
            callback(f"\n开始爬取政策详细内容，共 {len(all_policies)} 条政策")
        logging.info("\n" + "=" * 60)
        logging.info(f"▶▶ 开始爬取政策详细内容，共 {len(all_policies)} 条政策")
        logging.info("=" * 60)
        
        # 2. 爬取每个政策的详细内容
        for i, policy in enumerate(all_policies, 1):
            if self.stop_requested:
                if callback:
                    callback("停止爬取政策")
                logging.info("[停止] 停止爬取政策")
                break
            
            if callback:
                callback(f"\n进度: [{i}/{len(all_policies)}]")
            logging.info(f"\n进度: [{i}/{len(all_policies)}]")
            
            self.progress.current_policy_id = policy.id
            self.progress.current_policy_title = policy.title
            self._update_progress()
            
            success = self.crawl_single_policy(policy, callback)
            
            if success:
                self.progress.completed_count += 1
                self.progress.completed_policies.append(policy.id)
            else:
                self.progress.failed_count += 1
                self.progress.failed_policies.append({
                    'id': policy.id,
                    'title': policy.title,
                    'reason': '爬取失败'
                })
            
            self._update_progress()
            time.sleep(self.config.get("request_delay", 2))
        
        self.progress.end_time = datetime.now()
        self._update_progress()
        
        if callback:
            callback("\n爬取完成")
        logging.info("\n" + "=" * 60)
        logging.info("爬取完成")
        logging.info("=" * 60)
        logging.info(f"总计: {self.progress.total_count} 条")
        logging.info(f"成功: {self.progress.completed_count} 条")
        logging.info(f"失败: {self.progress.failed_count} 条")
        logging.info(f"成功率: {self.progress.success_rate:.2f}%")
        
        return self.progress
    
    def close(self):
        """关闭爬虫"""
        if hasattr(self.api_client, 'close'):
            self.api_client.close()
